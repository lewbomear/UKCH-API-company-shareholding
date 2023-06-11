import docx
import csv
import os
import json
import requests
from print_shareholder_info import print_shareholder_info
from dotenv import load_dotenv
from datetime import datetime


def generate_relevant_individual_info():
    """
    Generates information for an individual associate's company data
    """
    load_dotenv()

    document = docx.Document()

    API_KEY = os.environ.get("COMPANY_HOUSE_API_KEY")
    OFFICER_NAME = os.environ.get("OFFICER_NAME")
    OFFICER_DOB = os.environ.get("OFFICER_DOB")

    URL = f'https://api.company-information.service.gov.uk/search/officers?q="{OFFICER_NAME}"'

    headers = {"Authorization": API_KEY}

    # Add a title to the document
    document.add_heading(f"Associated companies for: {OFFICER_NAME}", level=1)
    current_heading = document.add_heading("Current appointments", level=1)
    current_end = document.add_paragraph("-")
    former_heading = document.add_heading("Former appointments", level=1)
    former_end = document.add_paragraph("--")

    # Load the SIC codes and descriptions from the CSV file
    with open(
        os.path.join(
            os.path.dirname(os.path.realpath(__file__)),
            "SIC07_CH_condensed_list_en.csv",
        ),
        newline="",
    ) as csvfile:
        sic_codes_reader = csv.reader(csvfile, delimiter=",", quotechar='"')
        sic_mapping = {row[0]: row[1] for row in sic_codes_reader}

    # Make the API request
    response = requests.get(URL, headers=headers)
    DATA = response.json()

    # only keep officers with exact name match
    exact_name_matches = []
    for item in DATA["items"]:
        if item["title"].lower() == OFFICER_NAME.lower():
            exact_name_matches.append(item)

    exact_name_dob_matches = []
    for item in DATA["items"]:
        if (
            item["title"].lower() == OFFICER_NAME.lower()
            and isinstance(item.get("date_of_birth"), dict)
            and str(item["date_of_birth"]["year"])
            + "-"
            + str(item["date_of_birth"]["month"]).zfill(2)
            == OFFICER_DOB[:7]
        ):
            exact_name_dob_matches.append(item)
    
    # print(json.dumps(exact_name_dob_matches, indent=4))
    print(len(exact_name_dob_matches))

    officer_data_cache = []  # Store data for all matches

    for match in exact_name_dob_matches:
        officer_data = []  # Store data for the current match
        page_no = 1
        items_per_page = 50
        start_index = 0
        remaining_results = 1
        total_results = -1
        multiple_pages = False

        link = match['links']['self']

        while remaining_results > 0:
            officer_url = f"https://api.company-information.service.gov.uk{link}?page={page_no}&items_per_page={items_per_page}&start_index={start_index}"
            officer_response = requests.get(officer_url, headers=headers)
            current_data = officer_response.json()  # gets the current data
            officer_data.append(current_data)  # adds it to the match data cache

            with open(f"{OFFICER_NAME} page {page_no}.json", "w", encoding="utf-8") as json_file:
                json.dump(current_data, json_file, ensure_ascii=False, indent=None)

            if total_results < 0:  # if you haven't checked the total results
                total_results = current_data["total_results"]  # fetch the number
                remaining_results = total_results  # set the remaining results to total
                multiple_pages = total_results > items_per_page

            remaining_results -= items_per_page
            start_index += items_per_page

            if remaining_results < items_per_page:
                items_per_page = remaining_results

            page_no += 1

        officer_data_cache.append(officer_data)

    if multiple_pages:
        merged_officer_data = []
        for data in officer_data_cache:
            merged_officer_data.extend(data)

        # Loop through the officer's appointments and print the company name, number, and nature of business
        for data in merged_officer_data:
            for appointment in data["items"]:
                company_name = appointment["appointed_to"]["company_name"]
                company_number = appointment["appointed_to"]["company_number"]
                appointed_on = appointment["appointed_on"]
                appointed_date = datetime.strptime(appointed_on, '%Y-%m-%d')
                formatted_appointed_date = appointed_date.strftime('%d %B %Y')
                officer_role = appointment["officer_role"]
                company_status = appointment["appointed_to"]["company_status"]
                psc_name = ""
                company_title = f"{company_name} ({company_number})"

                # Get the company profile URL
                company_profile_url = f"https://api.company-information.service.gov.uk/company/{company_number}"
                company_profile_response = requests.get(
                    company_profile_url, headers=headers
                )
                company_profile_data = company_profile_response.json()

                # Get the SIC code and look up the activity in the mapping dictionary
                sic_code = company_profile_data.get("sic_codes", ["N/A"])[0]
                activity = sic_mapping.get(sic_code, "Unknown")
            
                # Get company status, incorporation date, dissolution date

                company_inc = company_profile_data.get("date_of_creation")
                if "date_of_cessation" in company_profile_data:
                    company_dis = company_profile_data["date_of_cessation"]
                    dis_date = datetime.strptime(company_dis, '%Y-%m-%d')
                    formatted_dis_date = dis_date.strftime('%d %B %Y')
                else:
                    company_dis = None

                # Get the company's persons with significant control data
                company_psc_url = f"https://api.company-information.service.gov.uk/company/{company_number}/persons-with-significant-control"
                company_psc_response = requests.get(company_psc_url, headers=headers)
                company_psc_data = company_psc_response.json()

                # print(json.dumps(company_psc_data, indent=4))

                # Loop through the company's persons with significant control and print the name
                psc_names = []
                if "items" in company_psc_data and company_psc_data["items"]:
                    for psc in company_psc_data["items"]:
                        if "name" in psc:
                            psc_name = psc["name"]
                            psc_names.append(psc_name)
                else:
                    print("No PSC data available for this company.")

                if "resigned_on" in appointment:
                    resigned_on = appointment["resigned_on"]
                    resign_date = datetime.strptime(resigned_on, '%Y-%m-%d')
                    formatted_resign_date = resign_date.strftime('%d %B %Y')
                    if len(psc_names) == 1:
                        for psc_name in psc_names:
                            psc_statement = f"The company has a person with significant control named {psc_name}."
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement} \n"
                        )

                    elif len(psc_names) > 1:
                        last_name = psc_names.pop()
                        full_list = ", ".join(psc_names)
                        psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement} \n"
                        )

                    else:
                        psc_statement = (
                            "The company has no persons with significant control."
                        )
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement} \n"
                        )

                else:
                    if len(psc_names) == 1:
                        for psc_name in psc_names:
                            psc_statement = f"The company has a person with significant control named {psc_name}."
                            if "active" in company_status:
                                new_paragraph = current_end.insert_paragraph_before(
                                    f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} \n"
                                )
                            else:
                                new_paragraph = former_end.insert_paragraph_before(
                                    f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business was {activity}. {psc_statement} \n"
                                )
                    elif len(psc_names) > 1:
                        last_name = psc_names.pop()
                        full_list = ", ".join(psc_names)
                        psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                        if "active" in company_status:
                            new_paragraph = current_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} \n"
                            )
                        else:
                            new_paragraph = former_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business is {activity}. {psc_statement} \n"
                            )

                    else:
                        psc_statement = (
                            "The company has no persons with significant control."
                        )
                        if "active" in company_status:
                            new_paragraph = current_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} \n"
                            )
                        else:
                            new_paragraph = former_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business was {activity}. {psc_statement}\n"
                            )
    
    # Save the document as a Word file
    document.save(f"Associated companies for {OFFICER_NAME}.docx")
    document_text = docx.Document(f"Associated companies for {OFFICER_NAME}.docx")
    print(document_text)


if __name__ == "__main__":
    generate_relevant_individual_info()
