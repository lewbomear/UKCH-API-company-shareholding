from docx import Document
import requests
import csv
import json


document = Document()


api_key = "eb613ae5-33b1-4aaf-8f23-18ad86769709"


officer_name = "Keith Ranger DOLLIVER"
officer_dob = "1963-01"


url = (
    f'https://api.company-information.service.gov.uk/search/officers?q="{officer_name}"'
)


headers = {"Authorization": api_key}

# Add a title to the document
document.add_heading(f"Associated companies for: {officer_name}", level=1)
current_heading = document.add_heading("Current appointments", level=1)
current_end = document.add_paragraph("-")
former_heading = document.add_heading("Former appointments", level=1)
former_end = document.add_paragraph("--")

# Load the SIC codes and descriptions from the CSV file
with open("SIC07_CH_condensed_list_en.csv", newline="") as csvfile:
    sic_codes_reader = csv.reader(csvfile, delimiter=",", quotechar='"')
    sic_mapping = {row[0]: row[1] for row in sic_codes_reader}

# Make the API request
response = requests.get(url, headers=headers)
data = response.json()

# only keep officers with exact name match
exact_name_matches = []
for item in data["items"]:
    if item["title"].lower() == officer_name.lower():
        exact_name_matches.append(item)

exact_name_dob_matches = []
for item in data["items"]:
    if (
        item["title"].lower() == officer_name.lower()
        and isinstance(item.get("date_of_birth"), dict)
        and str(item["date_of_birth"]["year"])
        + "-"
        + str(item["date_of_birth"]["month"]).zfill(2)
        == officer_dob[:7]
    ):
        exact_name_dob_matches.append(item)

# Loop through the exact name matches and make a request to each URL in the links dictionary
for match in exact_name_dob_matches:
    officer_url = f"https://api.company-information.service.gov.uk{match['links']['self']}?page=1&items_per_page=100"
    officer_response = requests.get(officer_url, headers=headers)
    officer_data = officer_response.json()
    print(json.dumps(officer_data, indent=4))

    # Loop through the officer's appointments and print the company name, number, and nature of business
    for appointment in officer_data["items"]:
        company_name = appointment["appointed_to"]["company_name"]
        company_number = appointment["appointed_to"]["company_number"]
        appointed_on = appointment["appointed_on"]
        officer_role = appointment["officer_role"]
        company_status = appointment["appointed_to"]["company_status"]
        psc_name = ""
        company_title = f"{company_name} ({company_number})"

        # Get the company profile URL
        company_profile_url = (
            f"https://api.company-information.service.gov.uk/company/{company_number}"
        )
        company_profile_response = requests.get(company_profile_url, headers=headers)
        company_profile_data = company_profile_response.json()

        # Get the SIC code and look up the activity in the mapping dictionary
        sic_code = company_profile_data.get("sic_codes", ["N/A"])[0]
        activity = sic_mapping.get(sic_code, "Unknown")

        # Get company status, incorporation date, dissolution date

        company_inc = company_profile_data.get("date_of_creation")
        if "date_of_cessation" in company_profile_data:
            company_dis = company_profile_data["date_of_cessation"]
        else:
            company_dis = None

        # Get the company's persons with significant control data
        company_psc_url = f"https://api.company-information.service.gov.uk/company/{company_number}/persons-with-significant-control"
        company_psc_response = requests.get(company_psc_url, headers=headers)
        company_psc_data = company_psc_response.json()

        # print(json.dumps(company_psc_data, indent=4))

        # Loop through the company's persons with significant control and print the name
        psc_names = []
        if "items" in company_psc_data and company_psc_data["items"]:
            for psc in company_psc_data["items"]:
                if "name" in psc:
                    psc_name = psc["name"]
                    psc_names.append(psc_name)
        else:
            print("No PSC data available for this company.")

        if "resigned_on" in appointment:
            resigned_on = appointment["resigned_on"]
            if len(psc_names) == 1:
                for psc_name in psc_names:
                    psc_statement = f"The company has a person with significant control named {psc_name}."
                new_paragraph = former_end.insert_paragraph_before(
                    f"{company_name} ({company_number}) \n{officer_name} was appointed {officer_role} of {company_name} ({company_number}) on {appointed_on} and resigned on {resigned_on}. The nature of business is {activity}. {psc_statement}"
                )

            elif len(psc_names) > 1:
                last_name = psc_names.pop()
                full_list = ", ".join(psc_names)
                psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                new_paragraph = former_end.insert_paragraph_before(
                    f"{company_name} ({company_number}) \n{officer_name} was appointed {officer_role} of {company_name} ({company_number}) on {appointed_on} and resigned on {resigned_on}. The nature of business is {activity}. {psc_statement}"
                )

            else:
                psc_statement = "The company has no persons with significant control"
                new_paragraph = former_end.insert_paragraph_before(
                    f"{company_name} ({company_number}) \n{officer_name} was appointed {officer_role} of {company_name} on {appointed_on} and resigned on {resigned_on}. The nature of business is {activity}. {psc_statement}"
                )

        else:
            if len(psc_names) == 1:
                for psc_name in psc_names:
                    psc_statement = f"The company has a person with significant control named {psc_name}."
                    if "active" in company_status:
                        new_paragraph = current_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{officer_name} has been serving as {officer_role} of {company_name} since {appointed_on}. The nature of business is {activity}. {psc_statement}"
                        )
                    else:
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{officer_name} served as {officer_role} of {company_name} between {appointed_on} and {company_dis}. The nature of business was {activity}. {psc_statement}"
                        )
            elif len(psc_names) > 1:
                last_name = psc_names.pop()
                full_list = ", ".join(psc_names)
                psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                if "active" in company_status:
                    new_paragraph = current_end.insert_paragraph_before(
                        f"{company_name} ({company_number}) \n{officer_name} has been serving as {officer_role} of {company_name} since {appointed_on}. The nature of business is {activity}. {psc_statement}"
                    )
                else:
                    new_paragraph = former_end.insert_paragraph_before(
                        f"{company_name} ({company_number}) \n{officer_name} served as {officer_role} of {company_name} between {appointed_on} and {company_dis}. The nature of business is {activity}. {psc_statement}"
                    )

            else:
                psc_statement = "The company has no persons with significant control"
                if "active" in company_status:
                    new_paragraph = current_end.insert_paragraph_before(
                        f"{company_name} ({company_number}) \n{officer_name} has been serving as {officer_role} of {company_name} since {appointed_on}. The nature of business is {activity}. {psc_statement}"
                    )
                else:
                    new_paragraph = former_end.insert_paragraph_before(
                        f"{company_name} ({company_number}) \n{officer_name} served as {officer_role} of {company_name} between {appointed_on} and {company_dis}. The nature of business was {activity}. {psc_statement}"
                    )
# Save the document as a Word file
document.save(f"Associated companies for {officer_name}.docx")
