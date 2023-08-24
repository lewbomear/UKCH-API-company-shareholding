import docx
import csv
import os
import json
import requests
import openpyxl
import re
from pdf2image import convert_from_path
import pytesseract 
from PIL import Image
from dotenv import load_dotenv
from datetime import datetime
from openpyxl import Workbook
from audit_trail import save_page_as_pdf

def print_shareholder_info(company_number, company_name, sources_folder_path):
                
    # Insert API key
    API_KEY = os.environ.get("COMPANY_HOUSE_API_KEY")

    # Insert company number to retrieve its confirmation statement
    COMPANY_NUMBER = company_number

    # Make a filing history request to get all confirmation statements
    URL = f"https://api.company-information.service.gov.uk/company/{COMPANY_NUMBER}/filing-history?category=confirmation-statement&items_per_page=100"

    RESPONSE = requests.get(URL, auth=(API_KEY, ""))

    # Parse the response for the link to the document metadata
    response_json = RESPONSE.json()
    ITEMS = response_json["items"]

    # Check each confirmation statement to only select the latest confirmation statement with updates
    confirmation_statement_found = False
    for item in ITEMS:
        description = item["description"]
        if "confirmation-statement-with-updates" in description.lower():
            metadata_link = item["links"]["document_metadata"]
            metadata_response = requests.get(
                metadata_link,
                auth=(API_KEY, ""),
                headers={"Accept": "application/json"},
            )
            metadata_response_json = metadata_response.json()
            document_link = metadata_response_json["links"]["document"]

            # Get the content of the document
            document_response = requests.get(document_link, auth=(API_KEY, ""))
            document_content = document_response.content

            # Write the content to a file named "confirmation_statement.pdf"
            pdf_file_name = os.path.join(sources_folder_path, f"{company_name} confirmation statement.pdf")
            with open(pdf_file_name, "wb") as f:
                f.write(document_content)

            pages = convert_from_path(
                pdf_file_name,
                250,
                poppler_path=r"C:\Program Files\poppler-23.01.0\Library\bin",
            )

            images = []
            for page in pages:
                images.append(page)

            # Combine all the images into one
            combined_image = Image.new(
                "RGB", (images[0].width, sum([i.height for i in images]))
            )
            y_offset = 0
            for image in images:
                combined_image.paste(image, (0, y_offset))
                y_offset += image.height

            # Save the combined image (note for later: see if i can save the image as a temp file)
            combined_image.save("confirmation_statement.png", "PNG")

            # Set the path to Tesseract
            pytesseract.pytesseract.tesseract_cmd = (
                r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            )

            # Load the image
            combined_image = Image.open("confirmation_statement.png")

            # Use pytesseract to extract text from the image and format it into one paragraph, while limiting instances of three spaces to two
            text = (
                pytesseract.image_to_string(combined_image)
                .replace("\n\n", "\n")
                .replace("\n", "  ")
                .replace("confirmation  statement", "confirmation statement")
                .replace("this  confirmation", "this confirmation")
                .replace("of  this confirmation", "of this confirmation")
                .replace("date  of this", "date of this")
                .replace(r"\b([A-Z][a-z]+)  ([A-Z][a-z]+)\b", r"\1 \2")
            )

            
            # Check if the text contains "Full details of Shareholders"
            if "Full details of Shareholders" in text:
                confirmation_statement_found = True
                break

    if confirmation_statement_found:
        # Find all instances of "(?<!0 )(\d+) ORDINARY shares held as at the date of this confirmation statement" and the name that follows
        shares = re.findall(
            r"(?<!0 )(\d+) ORDINARY shares held as at the date of this confirmation statement  Name: (\S+(?: \S+)*)",
            text,
        )

        # Calculate the total number of shares
        total_shares = sum([int(share[0]) for share in shares])

        # Create a list of the shareholders and their percentage of shares
        shareholders = []
        for share in shares:
            percentage = round(int(share[0]) / total_shares * 100, 2)
            shareholders.append(f"- {share[1]} - {percentage}%")

        # Combine the shareholders into a string and print it
        shareholders_string = "\n".join(shareholders)
        ownership_statement = f"The company has the following shareholders:\n{shareholders_string}"
        
    else:
        ownership_statement = "No ownership information identified."
        

    return ownership_statement




def generate_relevant_individual_info():
    """
    Generates information for an individual associate's company data
    """
    load_dotenv()

    document = docx.Document()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.append(["Company", "Reg Number", "Status", "Officer Role", "Appointed On", "Resigned On", "Person with significant control"])
    worksheet.auto_filter.ref = worksheet.dimensions
    column_widths = [50, 15, 15, 15, 15, 15, 50]  

    for i, width in enumerate(column_widths, start=1):
        column_letter = openpyxl.utils.get_column_letter(i)
        worksheet.column_dimensions[column_letter].width = width

    
    API_KEY = os.environ.get("COMPANY_HOUSE_API_KEY")
    OFFICER_NAME = os.environ.get("OFFICER_NAME")
    OFFICER_DOB = os.environ.get("OFFICER_DOB")

    URL = f'https://api.company-information.service.gov.uk/search/officers?q="{OFFICER_NAME}"'

    headers = {"Authorization": API_KEY}

    # Add a title to the document
    document.add_heading(f"Associated companies for: {OFFICER_NAME}", level=1)
    current_heading = document.add_heading("Current appointments", level=1)
    current_end = document.add_paragraph("-")
    former_heading = document.add_heading("Former appointments", level=1)
    former_end = document.add_paragraph("--")

    # Create a folder with officer_name
    folder_path = os.path.join(".", OFFICER_NAME)  # Use current directory as the base path
    os.makedirs(folder_path, exist_ok=True)  # Create the folder if it doesn't exist
    

    # Create a 'sources' folder within the officer's folder
    sources_folder_path = os.path.join(folder_path, "sources")
    os.makedirs(sources_folder_path, exist_ok=True)

    # Load the SIC codes and descriptions from the CSV file
    with open(
        os.path.join(
            os.path.dirname(os.path.realpath(__file__)),
            "SIC07_CH_condensed_list_en.csv",
        ),
        newline="",
    ) as csvfile:
        sic_codes_reader = csv.reader(csvfile, delimiter=",", quotechar='"')
        sic_mapping = {row[0]: row[1] for row in sic_codes_reader}

    # Make the API request
    response = requests.get(URL, headers=headers)
    DATA = response.json()

    # only keep officers with exact name match
    exact_name_matches = []
    for item in DATA["items"]:
        if item["title"].lower() == OFFICER_NAME.lower():
            exact_name_matches.append(item)

    exact_name_dob_matches = []
    for item in DATA["items"]:
        if (
            item["title"].lower() == OFFICER_NAME.lower()
            and isinstance(item.get("date_of_birth"), dict)
            and str(item["date_of_birth"]["year"])
            + "-"
            + str(item["date_of_birth"]["month"]).zfill(2)
            == OFFICER_DOB[:7]
        ):
            exact_name_dob_matches.append(item)
    
    # Timer
    start_time = datetime.now()
    

    # print(json.dumps(exact_name_dob_matches, indent=4))
    print(len(exact_name_dob_matches))
    print("Generating...")
    officer_data_cache = []  # Store data for all matches

    for match in exact_name_dob_matches:
        officer_data = []  # Store data for the current match
        page_no = 1
        items_per_page = 50
        start_index = 0
        remaining_results = 1
        total_results = -1
        multiple_pages = False

        link = match['links']['self']

        while remaining_results > 0:
            officer_url = f"https://api.company-information.service.gov.uk{link}?page={page_no}&items_per_page={items_per_page}&start_index={start_index}"
            officer_response = requests.get(officer_url, headers=headers)
            current_data = officer_response.json()  # gets the current data
            officer_data.append(current_data)  # adds it to the match data cache

            with open(f"{OFFICER_NAME} page {page_no}.json", "w", encoding="utf-8") as json_file:
                json.dump(current_data, json_file, ensure_ascii=False, indent=None)

            if total_results < 0:  # if you haven't checked the total results
                total_results = current_data["total_results"]  # fetch the number
                remaining_results = total_results  # set the remaining results to total
                multiple_pages = total_results > items_per_page

            remaining_results -= items_per_page
            start_index += items_per_page

            if remaining_results < items_per_page:
                items_per_page = remaining_results

            page_no += 1

        officer_data_cache.append(officer_data)

        if multiple_pages:
            merged_officer_data = []
            for data in officer_data_cache:
                merged_officer_data.extend(data)
        else:
            merged_officer_data = []
            for data in officer_data_cache:
                merged_officer_data.extend(data)
            

        # Loop through the officer's appointments and print the company name, number, and nature of business
        for data in merged_officer_data:
            for appointment in data["items"]:
                company_name = appointment["appointed_to"]["company_name"]
                company_number = appointment["appointed_to"]["company_number"]
                appointed_on = appointment["appointed_on"]
                appointed_date = datetime.strptime(appointed_on, '%Y-%m-%d')
                formatted_appointed_date = appointed_date.strftime('%d %B %Y')
                officer_role = appointment["officer_role"]
                company_status = appointment["appointed_to"]["company_status"]
                psc_name = ""
                company_title = f"{company_name} ({company_number})"
                
                # Get the company profile URL
                company_profile_url = f"https://api.company-information.service.gov.uk/company/{company_number}"
                company_profile_response = requests.get(
                    company_profile_url, headers=headers
                )
                company_profile_data = company_profile_response.json()

                
                # Get the SIC code and look up the activity in the mapping dictionary
                sic_code = company_profile_data.get("sic_codes", ["N/A"])[0]
                activity = sic_mapping.get(sic_code, "Unknown")
            
                # Get company status, incorporation date, dissolution date

                company_inc = company_profile_data.get("date_of_creation")
                if "date_of_cessation" in company_profile_data:
                    company_dis = company_profile_data["date_of_cessation"]
                    dis_date = datetime.strptime(company_dis, '%Y-%m-%d')
                    formatted_dis_date = dis_date.strftime('%d %B %Y')
                else:
                    company_dis = None

                # Get the company's persons with significant control data
                company_psc_url = f"https://api.company-information.service.gov.uk/company/{company_number}/persons-with-significant-control"
                company_psc_response = requests.get(company_psc_url, headers=headers)
                company_psc_data = company_psc_response.json()

                # print(json.dumps(company_psc_data, indent=4))

                # Loop through the company's persons with significant control and print the name
                psc_names = []
                if "items" in company_psc_data and company_psc_data["items"]:
                    for psc in company_psc_data["items"]:
                        if "name" in psc:
                            psc_name = psc["name"]
                            psc_names.append(psc_name)
                else:
                    print(".")

                if "resigned_on" in appointment:
                    resigned_on = appointment["resigned_on"]
                    resign_date = datetime.strptime(resigned_on, '%Y-%m-%d')
                    formatted_resign_date = resign_date.strftime('%d %B %Y')
                    ownership_statement = print_shareholder_info(company_number, company_name, sources_folder_path)
                    if len(psc_names) == 1:
                        for psc_name in psc_names:
                            psc_statement = f"The company has a person with significant control named {psc_name}."
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                        )

                    elif len(psc_names) > 1:
                        last_name = psc_names.pop()
                        full_list = ", ".join(psc_names)
                        psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement}  {ownership_statement} \n"
                        )

                    else:
                        psc_statement = (
                            "The company has no persons with significant control."
                        )
                        new_paragraph = former_end.insert_paragraph_before(
                            f"{company_name} ({company_number}) \n{OFFICER_NAME} was appointed {officer_role} of {company_name} on {formatted_appointed_date} and resigned on {formatted_resign_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                        )

                else:
                    if len(psc_names) == 1:
                        ownership_statement = print_shareholder_info(company_number, company_name, sources_folder_path)
                        formatted_resign_date = "N/a"
                        for psc_name in psc_names:
                            psc_statement = f"The company has a person with significant control named {psc_name}."
                            if "active" in company_status:
                                new_paragraph = current_end.insert_paragraph_before(
                                    f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                                )
                            else:
                                new_paragraph = former_end.insert_paragraph_before(
                                    f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business was {activity}. {psc_statement} {ownership_statement} \n"
                                )
                    elif len(psc_names) > 1:
                        formatted_resign_date = "N/a"
                        last_name = psc_names.pop()
                        full_list = ", ".join(psc_names)
                        psc_statement = f"The company has the following persons with significant control: {full_list} and {last_name}."
                        ownership_statement = print_shareholder_info(company_number, company_name, sources_folder_path)
                        if "active" in company_status:
                            new_paragraph = current_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                            )
                        else:
                            new_paragraph = former_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                            )

                    else:
                        formatted_resign_date = "N/a"
                        ownership_statement = print_shareholder_info(company_number, company_name, sources_folder_path)
                        psc_statement = (
                            "The company has no persons with significant control."
                        )
                        if "active" in company_status:
                            formatted_resign_date = "N/a"
                            new_paragraph = current_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} has been serving as {officer_role} of {company_name} since {formatted_appointed_date}. The nature of business is {activity}. {psc_statement} {ownership_statement} \n"
                            )
                        else:
                            formatted_resign_date = "N/a"
                            new_paragraph = former_end.insert_paragraph_before(
                                f"{company_name} ({company_number}) \n{OFFICER_NAME} served as {officer_role} of {company_name} between {formatted_appointed_date} and {formatted_dis_date}. The nature of business was {activity}. {psc_statement} {ownership_statement} \n"
                            )
                worksheet.append([company_name, company_number, company_status, officer_role, formatted_appointed_date, formatted_resign_date, psc_name])
                
                pdf_url = f'https://find-and-update.company-information.service.gov.uk/company/{company_number}'
                pdf_path = fr'C:\Users\liubo\VSCode Projects\UKCH-API-company-shareholding-2\{OFFICER_NAME}\sources\{company_name}.pdf'
                save_page_as_pdf(pdf_url, pdf_path, company_name)
       
    
    end_time = datetime.now()
    duration = end_time - start_time

    # Calculate the duration
    duration = end_time - start_time

    # Calculate total minutes and seconds
    total_seconds = int(duration.total_seconds())
    minutes = total_seconds // 60
    seconds = total_seconds % 60
    

    # Save the document as a Word file
    document.save(os.path.join(folder_path, f"Associated companies for {OFFICER_NAME}.docx"))
    document_text = docx.Document(os.path.join(folder_path, f"Associated companies for {OFFICER_NAME}.docx"))
    print(document_text)
    workbook.save(os.path.join(folder_path, f"Associated companies for {OFFICER_NAME}.xlsx"))
    print("Completed!", f"Duration: {minutes} minutes {seconds} seconds")
    
 

    

if __name__ == "__main__":
    generate_relevant_individual_info()